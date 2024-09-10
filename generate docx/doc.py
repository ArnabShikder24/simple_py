from docx import Document

doc = Document()

doc.add_heading('Assignment: Payback Period Calculation', 0)

doc.add_paragraph('Course: SEISD (CSE 305)\nSession: Spring 2024\nSubmitted by: [Your Name]\nSubmission Date: 10th September 2024\n')

doc.add_heading('Objective:', level=1)
doc.add_paragraph('To calculate the payback period for a software system based on given costs and benefits.')

doc.add_heading('Costs and Benefits Breakdown:', level=1)
table = doc.add_table(rows=1, cols=3)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Cost / Benefit'
hdr_cells[1].text = 'Amount'
hdr_cells[2].text = 'Notes'

data = [
    ("Initial Investment", "", ""),
    ("Software License", "$50,000", "One-time cost"),
    ("Hardware Upgrades", "$10,000", "One-time cost"),
    ("Implementation Costs", "$20,000", "One-time cost"),
    ("Training Costs", "$5,000", "One-time cost"),
    ("Annual Recurring Costs", "", ""),
    ("Utilities (per year)", "$18,000", "$1,500 per month"),
    ("Maintenance & Support", "$10,000", "Annual cost"),
    ("Data Storage", "$2,000", "Annual cost"),
    ("Marketing Costs", "$40,000", "Conservative estimate"),
    ("Total Initial Costs (Year 1)", "$155,000", ""),
]

for cost, amount, notes in data:
    row_cells = table.add_row().cells
    row_cells[0].text = cost
    row_cells[1].text = amount
    row_cells[2].text = notes

doc.add_heading('Benefits:', level=1)
table = doc.add_table(rows=1, cols=3)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Benefit'
hdr_cells[1].text = 'Amount'
hdr_cells[2].text = 'Notes'

benefit_data = [
    ("Increased Sales (10% revenue boost)", "$100,000", "10% increase in $1M annual revenue"),
    ("Reduced Labor Costs", "$187,200", "Replacing 3 workers @ $30/hour"),
    ("Total Annual Benefits", "$287,200", ""),
]

for benefit, amount, notes in benefit_data:
    row_cells = table.add_row().cells
    row_cells[0].text = benefit
    row_cells[1].text = amount
    row_cells[2].text = notes

doc.add_heading('Payback Period Calculation:', level=1)
table = doc.add_table(rows=1, cols=4)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Year'
hdr_cells[1].text = 'Costs ($)'
hdr_cells[2].text = 'Benefits ($)'
hdr_cells[3].text = 'Net Gain/Loss ($)'

payback_data = [
    ("Year 1", "$155,000", "$287,200", "$132,200"),
    ("Year 2", "$30,000", "$244,120", "$214,120"),
    ("Year 3", "$30,000", "$207,502", "$177,502"),
]

for year, costs, benefits, net_gain in payback_data:
    row_cells = table.add_row().cells
    row_cells[0].text = year
    row_cells[1].text = costs
    row_cells[2].text = benefits
    row_cells[3].text = net_gain

doc.add_heading('Conclusion:', level=1)
doc.add_paragraph('The payback period is achieved within Year 1, as the net benefit in the first year exceeds the initial costs by $132,200. Any additional years increase the profitability further, even accounting for dollar depreciation.')

file_path = "D:\\Learning\\Automate Everything with Python\\generate docx\\Payback_Period_Assignment.docx"
doc.save(file_path)

file_path
